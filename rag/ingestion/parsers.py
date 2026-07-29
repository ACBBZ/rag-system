from __future__ import annotations

import io
from pathlib import Path
from typing import Literal, Protocol

import pandas as pd
from docx import Document
from PIL import Image
from pydantic import BaseModel, Field
from pypdf import PdfReader


class OcrCallable(Protocol):
    async def ocr(self, content: bytes, mime_type: str) -> str:
        ...


class ParsedBlock(BaseModel):
    block_type: Literal["heading", "paragraph", "list", "table", "code", "image_ocr"]
    text: str
    page: int | None = None
    title_path: list[str] = Field(default_factory=list)
    position: int
    metadata: dict[str, object] = Field(default_factory=dict)


class ParsedDocument(BaseModel):
    blocks: list[ParsedBlock]
    metadata: dict[str, object]

    @property
    def text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks if block.text.strip())


async def parse_document(
    filename: str,
    content: bytes,
    ocr_client: OcrCallable | None,
    *,
    max_pdf_pages: int = 1000,
    max_image_pixels: int = 40_000_000,
    max_spreadsheet_rows: int = 100_000,
) -> ParsedDocument:
    suffix = Path(filename).suffix.lower().lstrip(".")
    blocks: list[ParsedBlock] = []
    if suffix in {"txt", "md"}:
        text = content.decode("utf-8-sig")
        for position, section in enumerate(filter(str.strip, text.split("\n\n"))):
            blocks.append(
                ParsedBlock(
                    block_type="paragraph",
                    text=section.strip(),
                    position=position,
                    metadata={},
                )
            )
        return ParsedDocument(blocks=blocks, metadata={"file_type": suffix})

    if suffix == "pdf":
        reader = PdfReader(io.BytesIO(content))
        if len(reader.pages) > max_pdf_pages:
            raise ValueError("PDF page count exceeds configured limit")
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                blocks.append(
                    ParsedBlock(
                        block_type="paragraph",
                        text=text,
                        page=index,
                        position=len(blocks),
                        metadata={"page": index},
                    )
                )
        if not blocks and ocr_client is not None:
            text = await ocr_client.ocr(content, "application/pdf")
            blocks.append(
                ParsedBlock(
                    block_type="image_ocr",
                    text=text,
                    page=1,
                    position=0,
                    metadata={"ocr": True},
                )
            )
        return ParsedDocument(blocks=blocks, metadata={"file_type": suffix, "pages": len(reader.pages)})

    if suffix == "docx":
        doc = Document(io.BytesIO(content))
        title_path: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style.startswith("heading"):
                try:
                    level = max(1, int(style.split()[-1]))
                except ValueError:
                    level = 1
                title_path = [*title_path[: level - 1], text]
                block_type = "heading"
            else:
                block_type = "paragraph"
            blocks.append(
                ParsedBlock(
                    block_type=block_type,
                    text=text,
                    title_path=list(title_path),
                    position=len(blocks),
                    metadata={"style": style},
                )
            )
        for table_index, table in enumerate(doc.tables, start=1):
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            blocks.append(
                ParsedBlock(
                    block_type="table",
                    text="\n".join(rows),
                    title_path=list(title_path),
                    position=len(blocks),
                    metadata={"table": table_index, "rows": len(rows)},
                )
            )
        return ParsedDocument(blocks=blocks, metadata={"file_type": suffix, "tables": len(doc.tables)})

    if suffix == "csv":
        frame = pd.read_csv(io.BytesIO(content))
        return _frames_to_document({"csv": frame}, suffix, max_spreadsheet_rows)

    if suffix in {"xlsx", "xls"}:
        sheets = pd.read_excel(io.BytesIO(content), sheet_name=None)
        return _frames_to_document(sheets, suffix, max_spreadsheet_rows)

    if suffix in {"png", "jpg", "jpeg", "webp"}:
        with Image.open(io.BytesIO(content)) as image:
            width, height = image.size
            if width * height > max_image_pixels:
                raise ValueError("image pixel count exceeds configured limit")
            image.verify()
        if ocr_client is None:
            raise ValueError("OCR client is required for image parsing")
        text = await ocr_client.ocr(content, f"image/{'jpeg' if suffix == 'jpg' else suffix}")
        return ParsedDocument(
            blocks=[
                ParsedBlock(
                    block_type="image_ocr",
                    text=text,
                    page=1,
                    position=0,
                    metadata={"ocr": True},
                )
            ],
            metadata={"file_type": suffix},
        )
    raise ValueError(f"unsupported file type: {suffix}")


def _frames_to_document(
    frames: dict[str, pd.DataFrame],
    suffix: str,
    max_rows: int,
) -> ParsedDocument:
    blocks: list[ParsedBlock] = []
    for sheet_name, frame in frames.items():
        if len(frame.index) > max_rows:
            raise ValueError("spreadsheet row count exceeds configured limit")
        headers = [str(column) for column in frame.columns]
        window = 40
        for start in range(0, len(frame.index), window):
            subset = frame.iloc[start : start + window].fillna("")
            lines = [" | ".join(headers)]
            for row_index, (_, row) in enumerate(subset.iterrows(), start=start + 1):
                cells = ", ".join(
                    f"{column}: {row[column]}" for column in frame.columns
                )
                lines.append(f"row {row_index}: {cells}")
            blocks.append(
                ParsedBlock(
                    block_type="table",
                    text="\n".join(lines),
                    title_path=[str(sheet_name)],
                    position=len(blocks),
                    metadata={
                        "sheet": str(sheet_name),
                        "row_start": start + 1,
                        "row_end": start + len(subset.index),
                    },
                )
            )
    return ParsedDocument(blocks=blocks, metadata={"file_type": suffix, "sheets": list(frames)})
